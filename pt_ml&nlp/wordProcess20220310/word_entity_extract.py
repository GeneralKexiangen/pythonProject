import os
import docx
import json
import jieba.posseg as psg
import xlwt
import pandas as pd


def stopwordslist(stopwords_path):
    stopwords = [line.strip() for line in open(stopwords_path,encoding='UTF-8').readlines()]
    return stopwords

#常见停用词
sw = stopwordslist('stopword.txt')

# 自定义停用词
dsw = ['研末','血常规','局限','主诊','续服','油炸','患者','建议','食后','颈部','食用','右手','成功','过量','化验','炮','微弱','中服','成人',
'难治性','化学','出血点','处方','基础','证实','体质','总院','证候','复方','油漆','环境','微','盒','地方','时期','少','极易','饮用','图片','脚步',
'症状','原因','按语','生火','多用户','计数','生死考验','医疗费','吊水','教学','重','走路','月余','身体','过度','话题','理论','系统','全','右手','主诊',
'舌','现象','中药','县院','淡','药续用','附院','少见','天','细','建议','结论','单位','西药','姓名','医师','年龄','性别','病案','用药治疗','轻按','次数','体温',
'煎剂','皮肤','新','血液','手指','县院','血三系','原方']
dsw = list(set(dsw))

#分词并且去停用词
def cut_word(word):
    cws = psg.cut(word)
    finals = []
    for cw, flag in cws:
        if cw not in sw and flag in ('n','a','v','z') and cw not in dsw:
            finals.append(cw)
    return set(finals)

#需要处理的word文档路径
path = '/Users/zhiyue/Desktop/PT/NLP/血症紫癜出血贫血'
def parse_word():
    files = os.listdir(path)
    info_text = []
    for filename in files:
        portion = os.path.splitext(filename)
        if portion[1] == '.docx':
            print('>>>>>>>>>>>>>'+ filename)
            os.chdir(path)
            try:
                file = docx.Document(filename)
                meta_text_list = []
                trim_text_list = []
                for p in file.paragraphs:
                    print(p.text)
                    meta_text_list.append(p.text)
                    pro_text = str(p.text).strip()
                    pro_text = pro_text.replace('\n', '').replace('\t', '').replace('\r', '')
                    if len(pro_text) > 0:
                        trim_text_list.append(pro_text)
                info_text.append({'title':portion[0],'meta_text':meta_text_list,'trim_text':trim_text_list})
            except:
                print('未知异常')
        elif portion[1] == '' and portion[0][0:1] != '.':
            path1 = path +'/' +portion[0]+ '/'
            files1 = os.listdir(path1)
            for filename1 in files1:
                portion1 = os.path.splitext(filename1)
                if portion1[1] == '.docx':
                    print('>>>>>>>>>>>>>'+portion[0]+'/'+filename1)
                    os.chdir(path)
                    try:
                        file1 = docx.Document(path1+filename1)
                        meta_text_list = []
                        trim_text_list = []
                        for p in file1.paragraphs:
                            print(p.text)
                            meta_text_list.append(p.text)
                            pro_text = str(p.text).strip()
                            pro_text = pro_text.replace('\n','').replace('\t','').replace('\r','')
                            if len(pro_text)> 0:
                                trim_text_list.append(pro_text)
                        info_text.append(
                            {'title': portion[0]+ '/'+portion1[0], 'meta_text': meta_text_list, 'trim_text': trim_text_list})
                    except:
                        print('未知异常')
    print(json.dumps(info_text, ensure_ascii=False))
    return json.dumps(info_text, ensure_ascii=False)

info_text = parse_word()
df = pd.read_json(info_text,encoding="utf-8", orient='records')
df['sentences']= df['trim_text'].apply(lambda x :cut_word(str(x)))
df.to_excel('result.xls')
print('结果成功导出至 {0} 文件夹下'.format(path))